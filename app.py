import streamlit as st
import nltk
import string
import base64
import io
import random

from collections import Counter
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords

from pypdf import PdfReader
from docx import Document

import yake
from wordcloud import WordCloud
import plotly.express as px
import matplotlib.pyplot as plt
from gtts import gTTS

nltk.download('punkt', quiet=True)
nltk.download('punkt_tab', quiet=True)
nltk.download('stopwords', quiet=True)

st.set_page_config(
    page_title="NLP Document Summarizer",
    layout="wide",
    initial_sidebar_state="expanded"
)

if "dark_mode" not in st.session_state:
    st.session_state.dark_mode = False

def apply_css():
    if st.session_state.dark_mode:
        bg = "#121212"
        text_color = "#e0e0e0"
        card_bg = "rgba(255, 255, 255, 0.03)"
        border_color = "rgba(255, 255, 255, 0.08)"
        accent = "#ffffff"
        highlight_bg = "rgba(255, 255, 255, 0.1)"
        highlight_border = "rgba(255, 255, 255, 0.3)"
    else:
        bg = "#f8f9fa"
        text_color = "#212529"
        card_bg = "#ffffff"
        border_color = "rgba(0, 0, 0, 0.08)"
        accent = "#000000"
        highlight_bg = "rgba(0, 0, 0, 0.05)"
        highlight_border = "rgba(0, 0, 0, 0.2)"

    st.markdown(
        f"""
        <style>
        .stApp {{
            background: {bg};
            background-attachment: fixed;
            color: {text_color};
            font-family: 'Inter', 'Segoe UI', sans-serif;
        }}
        
        /* Clean Slider and Widget Styling */
        .stSlider > div > div > div {{
            padding-top: 5px !important;
            padding-bottom: 5px !important;
        }}
        .stSlider label {{
            font-size: 0.85rem !important;
            font-weight: 500 !important;
            color: {text_color} !important;
            opacity: 0.8;
            margin-bottom: -5px;
        }}
        div[data-testid="stSidebarUserContent"] hr {{
            margin-top: 1.5rem;
            margin-bottom: 1.5rem;
            border-color: {border_color};
            opacity: 0.5;
        }}
        .stToggle label p {{
            font-size: 0.9rem !important;
            font-weight: 500 !important;
            color: {text_color} !important;
        }}

        .glass-card {{
            background: {card_bg};
            border-radius: 8px;
            border: 1px solid {border_color};
            padding: 1.5rem;
            margin-bottom: 1.5rem;
            box-shadow: 0 2px 10px rgba(0, 0, 0, 0.02);
            transition: transform 0.2s ease-in-out;
        }}
        .metric-container {{
            display: flex;
            justify-content: space-between;
            gap: 1rem;
            margin-bottom: 1.5rem;
        }}
        .metric-card {{
            flex: 1;
            background: {card_bg};
            border-radius: 8px;
            border: 1px solid {border_color};
            padding: 1.5rem 1rem;
            text-align: center;
            box-shadow: 0 2px 10px rgba(0, 0, 0, 0.02);
        }}
        .metric-card h2 {{
            margin: 0;
            font-size: 2rem;
            font-weight: 600;
            color: {accent};
        }}
        .metric-card p {{
            margin: 0.5rem 0 0 0;
            opacity: 0.6;
            font-size: 0.85rem;
            text-transform: uppercase;
            letter-spacing: 0.5px;
            font-weight: 500;
        }}
        section[data-testid="stSidebar"] {{
            background: {card_bg};
            border-right: 1px solid {border_color};
        }}
        .stTabs [data-baseweb="tab-list"] {{
            gap: 2rem;
            background-color: transparent;
        }}
        .stTabs [data-baseweb="tab"] {{
            height: 3rem;
            white-space: pre-wrap;
            background-color: transparent;
            border-radius: 0;
            padding-top: 1rem;
            padding-bottom: 1rem;
            color: {text_color};
            opacity: 0.7;
        }}
        .stTabs [aria-selected="true"] {{
            border-bottom: 2px solid {accent} !important;
            color: {accent} !important;
            font-weight: 600;
            opacity: 1;
        }}
        h1, h2, h3, h4, h5, p, label, span, div {{
            color: {text_color};
        }}
        .highlight {{
            background-color: {highlight_bg};
            padding: 2px 4px;
            border-radius: 4px;
            border-bottom: 1px solid {highlight_border};
        }}
        .scrollable-content {{
            max-height: 500px;
            overflow-y: auto;
            padding-right: 10px;
            line-height: 1.6;
            font-size: 1rem;
        }}
        .scrollable-content::-webkit-scrollbar {{
            width: 4px;
        }}
        .scrollable-content::-webkit-scrollbar-thumb {{
            background-color: {border_color};
            border-radius: 4px;
        }}
        hr {{
            border-color: {border_color};
        }}
        </style>
        """,
        unsafe_allow_html=True
    )

apply_css()

with st.sidebar:
    st.markdown("<h2 style='text-align: center; margin-bottom: 0.5rem; font-weight: 600;'>NLP Document Summarizer</h2>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; opacity: 0.6; margin-bottom: 2rem; font-size: 0.85rem;'>Research Assistant Interface</p>", unsafe_allow_html=True)

    st.toggle("Dark Theme", key="dark_mode")
    st.markdown("---")
    
    uploaded_file = st.file_uploader("Upload Document", type=["pdf", "docx", "txt"])
    
    st.markdown("### Settings")
    summary_percent = st.slider("Summary Length (%)", 10, 90, 50)
    num_keywords = st.slider("Key Insights Count", 5, 20, 10)
    
    st.markdown("---")
    st.caption("System Status: Ready")

def extract_text(uploaded_file):
    if uploaded_file.type == "application/pdf":
        pdf = PdfReader(uploaded_file)
        return "".join([page.extract_text() or "" for page in pdf.pages])
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(uploaded_file)
        return "\n".join([para.text for para in doc.paragraphs])
    else:
        return str(uploaded_file.read(), "utf-8")

def extractive_summary(text, sentences, summary_percent):
    words = word_tokenize(text)
    stop_words = set(stopwords.words("english"))
    clean_words = [word.lower() for word in words if word.isalpha() and word.lower() not in stop_words]

    frequency = Counter(clean_words)
    if not frequency:
        return " ".join(sentences[:1]), set()

    max_frequency = max(frequency.values())
    for word in frequency:
        frequency[word] /= max_frequency

    sentence_scores = {}
    for sentence in sentences:
        for word in word_tokenize(sentence.lower()):
            if word in frequency:
                sentence_scores[sentence] = sentence_scores.get(sentence, 0) + frequency[word]

    ranked_sentences = sorted(sentence_scores.items(), key=lambda x: x[1], reverse=True)
    summary_length = max(1, int(len(sentences) * summary_percent / 100))
    top_sentences = ranked_sentences[:summary_length]
    selected = set(s for s, _ in top_sentences)

    return " ".join([s for s in sentences if s in selected]), selected

def extract_keywords(text, n=10):
    kw_extractor = yake.KeywordExtractor(top=n, n=2)
    return sorted(kw_extractor.extract_keywords(text), key=lambda x: x[1])

def color_func_gray(word, font_size, position, orientation, random_state=None, **kwargs):
    shade = random.randint(40, 100) if st.session_state.dark_mode else random.randint(20, 70)
    return f"hsl(0, 0%, {shade}%)"

def generate_wordcloud(text):
    wc = WordCloud(width=800, height=400, background_color=None, mode="RGBA", 
                   color_func=color_func_gray,
                   stopwords=set(stopwords.words("english"))).generate(text)
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.imshow(wc, interpolation="bilinear")
    ax.axis("off")
    fig.patch.set_alpha(0)
    return fig

def text_to_speech_bytes(text):
    buf = io.BytesIO()
    gTTS(text=text, lang="en").write_to_fp(buf)
    buf.seek(0)
    return buf

def make_docx_bytes(summary_text):
    doc = Document()
    doc.add_heading("SmartDoc AI - Summary", level=1)
    doc.add_paragraph(summary_text)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

if uploaded_file:
    text = extract_text(uploaded_file)
    if not text.strip():
        st.error("Error: Could not extract text from this document.")
        st.stop()

    words = word_tokenize(text)
    sentences = sent_tokenize(text)
    reading_time = max(1, round(len(words) / 200))
    avg_len = round(len(words) / len(sentences), 1) if sentences else 0

    st.markdown(f"""
        <div class="metric-container">
            <div class="metric-card"><h2>{len(words):,}</h2><p>Total Words</p></div>
            <div class="metric-card"><h2>{len(sentences):,}</h2><p>Sentences</p></div>
            <div class="metric-card"><h2>{avg_len}</h2><p>Avg Words / Sentence</p></div>
            <div class="metric-card"><h2>{reading_time} m</h2><p>Reading Time</p></div>
        </div>
    """, unsafe_allow_html=True)

    with st.spinner("Processing document data..."):
        summary, selected = extractive_summary(text, sentences, summary_percent)
        keywords = extract_keywords(text, n=num_keywords)
        summary_words_count = len(word_tokenize(summary))
        summary_reading_time = max(1, round(summary_words_count / 200))
        compression_ratio = (summary_words_count / len(words)) * 100 if words else 0

    tab1, tab2, tab3 = st.tabs(["Summary View", "Insights & Analytics", "Export & Audio"])

    with tab1:
        col_left, col_right = st.columns([1, 1], gap="large")
        with col_left:
            highlighted_html = "".join([f"<span class='highlight'>{s}</span> " if s in selected else f"{s} " for s in sentences])
            st.markdown(f"""
                <div class="glass-card">
                    <h4>Original Document</h4>
                    <hr/>
                    <div class="scrollable-content">{highlighted_html}</div>
                </div>
            """, unsafe_allow_html=True)
            
        with col_right:
            st.markdown(f"""
                <div class="glass-card">
                    <h4>Summarized</h4>
                    <hr/>
                    <div class="scrollable-content" style="font-weight: 400;">{summary}</div>
                </div>
            """, unsafe_allow_html=True)
            
            st.markdown(f"""
                <div class="glass-card" style="padding: 1.5rem; display: flex; justify-content: space-around; text-align: center;">
                    <div><h3 style="margin:0; font-weight:600;">{summary_reading_time} m</h3><p style="margin:0; opacity:0.6; font-size:0.8rem; text-transform:uppercase;">Adjusted Read Time</p></div>
                    <div><h3 style="margin:0; font-weight:600;">{compression_ratio:.1f}%</h3><p style="margin:0; opacity:0.6; font-size:0.8rem; text-transform:uppercase;">Compression Ratio</p></div>
                </div>
            """, unsafe_allow_html=True)

    with tab2:
        st.markdown('<div class="glass-card"><h4>Extracted Phrases</h4>', unsafe_allow_html=True)
        kw_cols = st.columns(3)
        for i, (kw, _) in enumerate(keywords):
            kw_cols[i % 3].markdown(f"<span style='opacity: 0.8;'>{i+1}.</span> {kw.title()}", unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

        chart_col1, chart_col2 = st.columns(2, gap="large")
        with chart_col1:
            st.markdown('<div class="glass-card"><h4>Frequency Map</h4>', unsafe_allow_html=True)
            st.pyplot(generate_wordcloud(text), use_container_width=True)
            st.markdown('</div>', unsafe_allow_html=True)

        with chart_col2:
            st.markdown('<div class="glass-card"><h4>Relevance Distribution</h4>', unsafe_allow_html=True)
            kw_labels, kw_scores = zip(*[(k, 1/(s+0.01)) for k, s in keywords[::-1]])
            
            fig = px.bar(x=kw_scores, y=kw_labels, orientation="h", color=kw_scores, color_continuous_scale="gray")
            fig.update_layout(
                paper_bgcolor="rgba(0,0,0,0)", 
                plot_bgcolor="rgba(0,0,0,0)", 
                showlegend=False, 
                xaxis_title="Relevance Score", 
                yaxis_title="", 
                font=dict(color="#e0e0e0" if st.session_state.dark_mode else "#212529"), 
                margin=dict(l=0, r=0, t=0, b=0), 
                coloraxis_showscale=False
            )
            st.plotly_chart(fig, use_container_width=True)
            st.markdown('</div>', unsafe_allow_html=True)

    with tab3:
        st.markdown('<div class="glass-card"><h4>Audio Synthesis</h4><p style="opacity: 0.7;">Generate playback for the synthesized summary.</p>', unsafe_allow_html=True)
        if st.button("Initialize Audio Generation", use_container_width=True):
            with st.spinner("Synthesizing audio data..."):
                st.audio(text_to_speech_bytes(summary), format="audio/mp3")
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="glass-card"><h4>Export Data</h4><p style="opacity: 0.7;">Download structured data for local use.</p>', unsafe_allow_html=True)
        e1, e2 = st.columns(2)
        e1.download_button("Export as Plain Text (.txt)", summary, "summary.txt", "text/plain", use_container_width=True)
        e2.download_button("Export as Document (.docx)", make_docx_bytes(summary), "summary.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
        st.markdown('</div>', unsafe_allow_html=True)

else:
    st.markdown("""
        <div style="display: flex; flex-direction: column; align-items: center; justify-content: center; height: 60vh; text-align: center; opacity: 0.8;">
            <h2 style="font-weight: 600; margin-bottom: 1rem;">NLP Document Summarizer</h2>
            <p style="max-width: 500px; font-size: 1rem; opacity: 0.7;">
                Upload a standard document format (PDF, DOCX, TXT) via the sidebar controls to initialize processing, extraction, and analytics modules.
            </p>
        </div>
    """, unsafe_allow_html=True)
